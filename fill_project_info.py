#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
根据"项目信息.txt"中的配置，替换"标准化模板"目录下所有docx文件中的{字段}标记，结果保存到"输出文件"目录
"""
import os
import shutil
import re
import configparser
from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import load_workbook
from openpyxl.utils.exceptions import InvalidFileException

# 读取配置文件
config = configparser.ConfigParser()
config_path = "config.ini"
if os.path.exists(config_path):
    config.read(config_path, encoding="utf-8")

# 配置，优先使用配置文件中的值，否则使用默认值
TEMPLATE_DIR = config.get("paths", "template_dir", fallback="标准化模板")
OUTPUT_DIR = config.get("paths", "output_dir", fallback="输出文件")
PROJECT_INFO_FILE = config.get("paths", "project_info_file", fallback="项目信息.txt")

# 收集模板中未定义的标签
undefined_tags = set()

def load_project_info():
    """加载项目信息，返回字段字典。支持单行(key:value)和多行(key:```...```)格式"""
    info = {}
    if not os.path.exists(PROJECT_INFO_FILE):
        print(f"错误：{PROJECT_INFO_FILE} 文件不存在！")
        return None

    try:
        with open(PROJECT_INFO_FILE, "r", encoding="utf-8") as f:
            lines = f.readlines()

        i = 0
        while i < len(lines):
            line = lines[i].rstrip("\n\r")
            stripped = line.strip()

            # 跳过空行和注释
            if not stripped or stripped.startswith("#"):
                i += 1
                continue

            # 支持英文冒号(:)和中文冒号(：)
            sep = ":" if ":" in stripped else ("：" if "：" in stripped else None)
            if sep:
                key, value = stripped.split(sep, 1)
                key = key.strip()
                value = value.strip()

                # 检测多行块：key:``` 或 key: ``` 或 key:\n```
                if value == "```":
                    # 读取直到结束的 ```
                    multiline_parts = []
                    i += 1
                    while i < len(lines):
                        ml = lines[i].rstrip("\n\r")
                        if ml.strip() == "```":
                            break
                        multiline_parts.append(ml)
                        i += 1
                    info[key] = "\n".join(multiline_parts)
                elif value == "" and i + 1 < len(lines) and lines[i + 1].strip() == "```":
                    # key: (空值) 且下一行是 ```，视为多行块
                    i += 2
                    multiline_parts = []
                    while i < len(lines):
                        ml = lines[i].rstrip("\n\r")
                        if ml.strip() == "```":
                            break
                        multiline_parts.append(ml)
                        i += 1
                    info[key] = "\n".join(multiline_parts)
                else:
                    info[key] = value

            i += 1

        print(f"成功加载 {len(info)} 个项目字段")
        return info
    except Exception as e:
        print(f"读取 {PROJECT_INFO_FILE} 失败: {str(e)}")
        return None

def _check_undefined_tags(text, info, source_desc):
    """检查文本中的{标签}是否在info中存在定义，未定义则警告"""
    matches = re.findall(r'\{([^}]+)\}', text)
    for tag in matches:
        if tag not in info:
            undefined_tags.add(tag)
            print(f"  警告：标签 {{{tag}}} 在项目信息中未定义！ (来源: {source_desc})")

def _add_line_break(run):
    """在run中插入一个换行符(w:br)"""
    br = OxmlElement("w:br")
    run._element.append(br)

def _replace_in_paragraph(para, info, source_desc=""):
    """在段落run级别替换字段，保留原格式，支持多行内容（\n渲染为换行）"""
    full_text = para.text
    if "{" not in full_text:
        return

    # 检查未定义标签
    _check_undefined_tags(full_text, info, source_desc)

    # 检查是否有需要替换的内容
    has_placeholder = False
    for key in info:
        if "{" + key + "}" in full_text:
            has_placeholder = True
            break
    if not has_placeholder:
        return

    runs = para.runs
    if not runs:
        return

    # 在拼接文本上执行替换
    new_text = full_text
    for key, value in info.items():
        placeholder = "{" + key + "}"
        new_text = new_text.replace(placeholder, value)

    if new_text == full_text:
        return

    # 检查替换结果是否包含换行
    if "\n" not in new_text:
        # 无换行，按原策略处理
        runs[0].text = new_text
        for run in runs[1:]:
            run.text = ""
    else:
        # 有换行，需要在run中插入<w:br/>元素
        # 清空所有run
        for run in runs:
            run.text = ""

        # 将文本按\n分割，逐段写入第一个run，每段之间插入<w:br/>
        parts = new_text.split("\n")
        for idx, part in enumerate(parts):
            if idx > 0:
                _add_line_break(runs[0])
            runs[0].text = runs[0].text + part

def replace_text_in_doc(doc, info, source_desc=""):
    """替换文档中的所有字段，保留原格式"""
    # 替换段落中的字段
    for para in doc.paragraphs:
        _replace_in_paragraph(para, info, source_desc)

    # 替换表格中的字段
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, info, source_desc)

    return doc

def replace_text_in_excel(wb, info, source_desc=""):
    """替换Excel文件中的所有字段，支持复合内容和合并单元格"""
    for ws in wb.worksheets:
        # 获取合并单元格信息
        merged_ranges = list(ws.merged_cells.ranges)

        # 收集合并区域中左上角单元格的坐标
        merge_top_left = set()
        for merged_range in merged_ranges:
            merge_top_left.add((merged_range.min_row, merged_range.min_col))

        # 遍历所有单元格
        for row in ws.iter_rows():
            for cell in row:
                # 合并单元格：只有左上角单元格有值，跳过其他合并区域的单元格
                if (cell.row, cell.column) not in merge_top_left:
                    # 检查是否在某个合并区域内
                    in_merged = False
                    for merged_range in merged_ranges:
                        if (merged_range.min_row <= cell.row <= merged_range.max_row and
                            merged_range.min_col <= cell.column <= merged_range.max_col):
                            if cell.row != merged_range.min_row or cell.column != merged_range.min_col:
                                in_merged = True
                                break
                    if in_merged:
                        continue

                cell_value = cell.value
                if cell_value is None:
                    continue

                # 支持字符串类型的复合内容替换
                if isinstance(cell_value, str):
                    if "{" not in cell_value:
                        continue

                    # 检查未定义标签
                    _check_undefined_tags(cell_value, info, source_desc)

                    new_value = cell_value
                    for key, value in info.items():
                        placeholder = "{" + key + "}"
                        new_value = new_value.replace(placeholder, value)

                    if new_value != cell_value:
                        cell.value = new_value

    return wb

def process_docx_file(source_path, target_path, info):
    """处理单个docx文件，替换字段后保存到目标路径"""
    try:
        # 跳过临时文件和损坏文件
        if os.path.basename(source_path).startswith("~$"):
            print(f"跳过临时文件: {source_path}")
            return False
        doc = Document(source_path)
        doc = replace_text_in_doc(doc, info, source_path)
        doc.save(target_path)
        return True
    except Exception as e:
        print(f"处理文件失败 {source_path}: {str(e)}")
        return False

def process_excel_file(source_path, target_path, info):
    """处理单个Excel文件，替换字段后保存到目标路径"""
    try:
        # 跳过临时文件
        if os.path.basename(source_path).startswith("~$"):
            print(f"跳过临时文件: {source_path}")
            return False
        # 加载Excel文件，keep_vba=True保留宏（支持.xlsm）
        wb = load_workbook(source_path, keep_vba=True)
        wb = replace_text_in_excel(wb, info, source_path)
        wb.save(target_path)
        return True
    except InvalidFileException as e:
        print(f"无效Excel文件 {source_path}: {str(e)}")
        return False
    except Exception as e:
        print(f"处理Excel文件失败 {source_path}: {str(e)}")
        return False

def main():
    # 加载项目信息
    info = load_project_info()
    if not info:
        return

    # 检查模板目录
    if not os.path.exists(TEMPLATE_DIR):
        print(f"错误：{TEMPLATE_DIR} 目录不存在！请先运行 doc_to_docx.py 生成标准化模板")
        return

    # 创建输出目录
    Path(OUTPUT_DIR).mkdir(parents=True, exist_ok=True)

    # 统计信息
    total_files = 0
    processed_docx = 0
    processed_excel = 0
    copied_files = 0
    failed_files = 0

    # 遍历模板目录
    for root, dirs, files in os.walk(TEMPLATE_DIR):
        # 计算相对路径
        rel_path = os.path.relpath(root, TEMPLATE_DIR)
        target_root = os.path.join(OUTPUT_DIR, rel_path)

        # 创建输出子目录
        Path(target_root).mkdir(parents=True, exist_ok=True)

        for file in files:
            total_files += 1
            source_file = os.path.join(root, file)
            target_file = os.path.join(target_root, file)

            # 跳过临时文件
            if file.startswith("~$"):
                print(f"跳过临时文件: {source_file}")
                continue

            # 处理docx文件
            if file.lower().endswith(".docx"):
                print(f"处理Word: {source_file} -> {target_file}")
                if process_docx_file(source_file, target_file, info):
                    processed_docx += 1
                else:
                    failed_files += 1
                    # 处理失败的话复制原文件
                    print(f"复制原文件: {source_file} -> {target_file}")
                    shutil.copy2(source_file, target_file)
            # 处理Excel文件
            elif file.lower().endswith((".xlsx", ".xlsm")):
                print(f"处理Excel: {source_file} -> {target_file}")
                if process_excel_file(source_file, target_file, info):
                    processed_excel += 1
                else:
                    failed_files += 1
                    # 处理失败的话复制原文件
                    print(f"复制原文件: {source_file} -> {target_file}")
                    shutil.copy2(source_file, target_file)
            else:
                # 其他文件直接复制
                print(f"复制: {source_file} -> {target_file}")
                shutil.copy2(source_file, target_file)
                copied_files += 1

    # 输出统计
    print(f"\n处理完成！")
    print(f"总文件数: {total_files}")
    print(f"成功处理Word文件: {processed_docx}")
    print(f"成功处理Excel文件: {processed_excel}")
    print(f"直接复制其他文件: {copied_files}")
    print(f"处理失败文件: {failed_files}")

    # 输出未定义标签汇总
    if undefined_tags:
        print(f"\n⚠ 发现 {len(undefined_tags)} 个未定义标签（在项目信息中找不到对应字段）：")
        for tag in sorted(undefined_tags):
            print(f"  - {{{tag}}}")

    print(f"结果已保存到: {OUTPUT_DIR} 目录")

    # 将输出目录中所有 .xlsx 文件重命名为 .xls
    renamed = 0
    for root, _, files in os.walk(OUTPUT_DIR):
        for file in files:
            if file.lower().endswith(".xlsx"):
                old_path = os.path.join(root, file)
                new_path = os.path.join(root, file[:-5] + ".xls")
                os.rename(old_path, new_path)
                print(f"重命名: {old_path} -> {new_path}")
                renamed += 1
    if renamed:
        print(f"已将 {renamed} 个 .xlsx 文件重命名为 .xls")

if __name__ == "__main__":
    main()
    print("\n按任意键退出...")
    input()
